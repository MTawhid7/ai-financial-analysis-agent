"""Word (.docx) exporter — converts a Markdown report to a styled Word document."""

from __future__ import annotations

import re
from pathlib import Path


def export_to_docx(markdown_text: str, output_path: Path) -> Path:
    """Convert markdown_text to a .docx file and write to output_path."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise ImportError(
            "Word export requires python-docx. Install with: pip install python-docx"
        ) from exc

    doc = Document()

    # Set document font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 60)
        elif stripped.startswith(("- ", "* ", "+ ")):
            # Bullet list item — strip inline markdown before adding
            text = _strip_inline(stripped[2:])
            doc.add_paragraph(text, style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            text = _strip_inline(re.sub(r"^\d+\. ", "", stripped))
            doc.add_paragraph(text, style="List Number")
        elif stripped.startswith(">"):
            p = doc.add_paragraph(stripped.lstrip("> "))
            p.paragraph_format.left_indent = Pt(24)
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, stripped)

    doc.save(str(output_path))
    return output_path


def _strip_inline(text: str) -> str:
    """Remove markdown inline syntax (bold, italic, code, links)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _add_inline_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, applying bold/italic/code formatting."""
    # Split on bold (**...**), italic (*...*), and code (`...`)
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    from docx.shared import Pt, RGBColor
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.bold = True
            run.text = part[2:-2]
        elif part.startswith("*") and part.endswith("*"):
            run.italic = True
            run.text = part[1:-1]
        elif part.startswith("`") and part.endswith("`"):
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            run.text = part[1:-1]
        else:
            run.text = part
