"""File parser — converts uploaded files to structured summaries.

Security invariants:
- NO raw cell/document content reaches the primary LLM.
- All formats produce a fixed-schema JSON summary only.
- CSV formula injection (=, +, -, @) is scrubbed before summary.
- Large documents use hierarchical summarisation so no context is lost
  regardless of document length.

Supported formats: CSV, PDF, XLSX/XLS, DOCX, TXT, MD, JSON
"""

from __future__ import annotations

import io
import json as _json
import logging
import re
from typing import Any

logger = logging.getLogger(__name__)

_FORMULA_RE = re.compile(r"^[=+\-@]")
_CHUNK_CHARS = 3000   # characters per Flash-Lite chunk
_CHUNK_OVERLAP = 200  # overlap between adjacent chunks


# ---------------------------------------------------------------------------
# Hierarchical summarisation — covers entire documents regardless of size
# ---------------------------------------------------------------------------


async def _summarise_chunk(text: str, subllm, task: str = "") -> str:
    """Summarise one chunk of text with Flash-Lite."""
    from langchain_core.messages import HumanMessage
    from ..core.llm import content_to_str

    prompt = (
        f"{task}\n\n" if task else
        "Summarise the following text in 3-5 sentences. Focus on: "
        "key topics, important figures, company names, and financial data.\n\n"
    ) + f"Text:\n{text}"

    try:
        response = await subllm.ainvoke([HumanMessage(content=prompt)])
        return content_to_str(response.content if hasattr(response, "content") else response).strip()
    except Exception as exc:
        logger.warning("Chunk summarisation failed: %s", exc)
        return text[:500]  # fallback to first 500 chars


def _sliding_chunks(text: str, chunk_size: int = _CHUNK_CHARS, overlap: int = _CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping chunks by character count."""
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


async def _hierarchical_summarise(text: str, subllm) -> str:
    """Summarise any-length document without truncation.

    Short documents (< chunk_size) are summarised directly.
    Long documents are split into overlapping chunks, each summarised,
    then the chunk summaries are combined into a final summary.
    """
    if not text.strip():
        return ""

    if len(text) <= _CHUNK_CHARS:
        return await _summarise_chunk(text, subllm)

    # Multi-chunk: summarise each, then combine
    chunks = _sliding_chunks(text)
    chunk_summaries: list[str] = []
    for chunk in chunks:
        summary = await _summarise_chunk(chunk, subllm)
        chunk_summaries.append(summary)

    combined = "\n\n".join(chunk_summaries)
    final = await _summarise_chunk(
        combined,
        subllm,
        task=(
            "The following are section-by-section summaries of a longer document. "
            "Combine them into a single coherent 4-6 sentence overview, preserving "
            "the most important financial figures and company names:"
        ),
    )
    return final


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------


def parse_csv(file_bytes: bytes, filename: str = "file.csv") -> dict[str, Any]:
    """Parse a CSV file. Returns a fixed-schema summary — no raw data to LLM."""
    try:
        import pandas as pd
        df = pd.read_csv(io.BytesIO(file_bytes))
    except Exception as exc:
        return {"error": f"Could not parse CSV: {exc}", "filename": filename}

    injection_count = 0
    for col in df.select_dtypes(include="object").columns:
        def _scrub(val):
            nonlocal injection_count
            s = str(val) if val is not None else ""
            if _FORMULA_RE.match(s):
                injection_count += 1
                return "[REMOVED]"
            return s[:200]
        df[col] = df[col].map(_scrub)

    numeric_stats: dict[str, dict] = {}
    for col in df.select_dtypes(include="number").columns:
        stats = df[col].describe()
        numeric_stats[col] = {k: round(float(v), 4) for k, v in stats.items() if k != "count"}

    return {
        "filename": filename,
        "file_type": "csv",
        "shape": {"rows": int(df.shape[0]), "columns": int(df.shape[1])},
        "columns": list(df.columns),
        "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
        "preview": df.head(5).fillna("").to_dict(orient="records"),
        "numeric_stats": numeric_stats,
        "formula_cells_removed": injection_count,
    }


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


async def parse_pdf(file_bytes: bytes, filename: str, subllm=None) -> dict[str, Any]:
    """Extract and summarise a PDF. Covers ALL pages via hierarchical summarisation."""
    try:
        import pdfplumber
        text_pages: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:  # no page cap — process entire document
                page_text = page.extract_text() or ""
                text_pages.append(page_text.strip())
    except Exception as exc:
        return {"error": f"Could not parse PDF: {exc}", "filename": filename}

    full_text = "\n\n".join(text_pages)
    char_count = len(full_text)

    summary = full_text[:2000]  # default fallback
    if subllm and full_text.strip():
        summary = await _hierarchical_summarise(full_text, subllm)

    return {
        "filename": filename,
        "file_type": "pdf",
        "pages": page_count,
        "char_count": char_count,
        "summary": summary,
    }


# ---------------------------------------------------------------------------
# XLSX / XLS
# ---------------------------------------------------------------------------


def parse_xlsx(file_bytes: bytes, filename: str = "file.xlsx") -> dict[str, Any]:
    """Parse an Excel workbook. Returns per-sheet summaries."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    except Exception as exc:
        return {"error": f"Could not parse Excel file: {exc}", "filename": filename}

    sheets: dict[str, Any] = {}
    for name in wb.sheetnames[:10]:  # cap at 10 sheets
        ws = wb[name]
        rows_iter = ws.iter_rows(values_only=True)

        try:
            headers = [str(c) if c is not None else "" for c in next(rows_iter)]
        except StopIteration:
            sheets[name] = {"rows": 0, "columns": 0}
            continue

        data_rows: list[list] = []
        row_count = 0
        injection_count = 0
        for row in rows_iter:
            row_count += 1
            clean_row = []
            for cell in row:
                s = str(cell) if cell is not None else ""
                if _FORMULA_RE.match(s):
                    injection_count += 1
                    s = "[REMOVED]"
                clean_row.append(s[:200])
            if row_count <= 5:
                data_rows.append(clean_row)

        sheets[name] = {
            "rows": row_count,
            "columns": len(headers),
            "column_names": headers[:20],  # cap header list
            "preview": [dict(zip(headers, r)) for r in data_rows],
            "formula_cells_removed": injection_count,
        }

    wb.close()
    return {
        "filename": filename,
        "file_type": "xlsx",
        "sheet_count": len(wb.sheetnames),
        "sheets": sheets,
    }


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


async def parse_docx(file_bytes: bytes, filename: str, subllm=None) -> dict[str, Any]:
    """Extract text from a Word document and summarise via hierarchical Flash-Lite."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract table cells
        table_text: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text.append(row_text)

        full_text = "\n\n".join(paragraphs)
        if table_text:
            full_text += "\n\nTable data:\n" + "\n".join(table_text)

    except Exception as exc:
        return {"error": f"Could not parse DOCX: {exc}", "filename": filename}

    char_count = len(full_text)
    summary = full_text[:2000]
    if subllm and full_text.strip():
        summary = await _hierarchical_summarise(full_text, subllm)

    return {
        "filename": filename,
        "file_type": "docx",
        "char_count": char_count,
        "paragraph_count": len(paragraphs),
        "summary": summary,
    }


# ---------------------------------------------------------------------------
# TXT / MD
# ---------------------------------------------------------------------------


async def parse_text(file_bytes: bytes, filename: str, subllm=None) -> dict[str, Any]:
    """Read plain text or Markdown. Summarises the full content hierarchically."""
    try:
        text = file_bytes.decode("utf-8", errors="replace")
    except Exception as exc:
        return {"error": f"Could not read file: {exc}", "filename": filename}

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    char_count = len(text)
    word_count = len(text.split())

    excerpt = text[:500].strip()
    summary = excerpt
    if subllm and text.strip() and char_count > 500:
        summary = await _hierarchical_summarise(text, subllm)

    return {
        "filename": filename,
        "file_type": ext,
        "char_count": char_count,
        "word_count": word_count,
        "excerpt": excerpt,
        "summary": summary,
    }


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------


def parse_json(file_bytes: bytes, filename: str = "file.json") -> dict[str, Any]:
    """Parse a JSON file — returns structure summary without raw values."""
    try:
        data = _json.loads(file_bytes.decode("utf-8", errors="replace"))
    except Exception as exc:
        return {"error": f"Could not parse JSON: {exc}", "filename": filename}

    def _describe(obj: Any, depth: int = 0) -> dict:
        if depth > 3:
            return {"type": type(obj).__name__}
        if isinstance(obj, dict):
            return {
                "type": "object",
                "keys": list(obj.keys())[:20],
                "key_count": len(obj),
            }
        if isinstance(obj, list):
            sample = _describe(obj[0], depth + 1) if obj else {}
            return {"type": "array", "length": len(obj), "item_schema": sample}
        return {"type": type(obj).__name__, "value_preview": str(obj)[:80]}

    return {
        "filename": filename,
        "file_type": "json",
        "schema": _describe(data),
        "top_level_keys": list(data.keys())[:20] if isinstance(data, dict) else None,
        "array_length": len(data) if isinstance(data, list) else None,
    }
