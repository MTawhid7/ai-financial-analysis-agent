"""Word (DOCX) file parser."""
from __future__ import annotations

import io
from typing import Any


async def parse_docx(file_bytes: bytes, filename: str, subllm=None) -> dict[str, Any]:
    """Extract text and table content from a Word document, then summarise."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs  = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                line = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if line:
                    table_lines.append(line)
        full_text = "\n\n".join(paragraphs)
        if table_lines:
            full_text += "\n\nTable data:\n" + "\n".join(table_lines)
    except Exception as exc:
        return {"error": f"Could not parse DOCX: {exc}", "filename": filename}

    char_count = len(full_text)
    summary    = full_text[:2000]
    if subllm and full_text.strip():
        from ._summarise import hierarchical_summarise
        summary = await hierarchical_summarise(full_text, subllm)

    return {
        "filename":        filename,
        "file_type":       "docx",
        "char_count":      char_count,
        "paragraph_count": len(paragraphs),
        "summary":         summary,
    }
