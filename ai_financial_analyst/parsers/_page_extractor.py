"""Structured page/section extraction from all supported document formats.

Each format produces a list of RawPage objects — the atomic input to the
PageIndex indexing pipeline. The existing summary parsers remain unchanged;
this module is a parallel code path used exclusively by the background indexer.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any


@dataclass
class RawPage:
    """One physical page (PDF) or one logical section (DOCX, MD, HTML, TXT)."""
    page_number: int
    content: str
    heading_breadcrumb: list[str] = field(default_factory=list)
    section_path: str = ""
    tables: list[dict[str, Any]] = field(default_factory=list)  # [{"headers":[…], "rows":[[…]]}]
    has_figures: bool = False
    is_toc: bool = False
    is_bibliography: bool = False

    @property
    def word_count(self) -> int:
        return len(self.content.split())

    @property
    def token_estimate(self) -> int:
        return max(1, len(self.content) // 4)  # rough 4 chars/token heuristic


# ---------------------------------------------------------------------------
# PDF extractor (native text + table detection)
# ---------------------------------------------------------------------------

def extract_pdf_pages(file_bytes: bytes) -> list[RawPage]:
    """Extract structured pages from a native (text-selectable) PDF.

    Falls back to empty pages for scanned PDFs — the OCR module handles those.
    """
    try:
        import pdfplumber
    except ImportError:
        return []

    pages: list[RawPage] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            tables = _extract_pdfplumber_tables(page)
            is_toc = _looks_like_toc(text)
            is_bib = _looks_like_bibliography(text)
            heading_breadcrumb = _detect_headings_from_text(text)
            pages.append(RawPage(
                page_number=i,
                content=text,
                heading_breadcrumb=heading_breadcrumb,
                section_path=str(i),
                tables=tables,
                has_figures=_has_figure_markers(text),
                is_toc=is_toc,
                is_bibliography=is_bib,
            ))
    return pages


def _extract_pdfplumber_tables(page: Any) -> list[dict]:
    """Extract tables from a pdfplumber page object."""
    tables = []
    try:
        for raw_table in page.extract_tables() or []:
            if not raw_table:
                continue
            headers = [str(c or "").strip() for c in raw_table[0]]
            rows = []
            for row in raw_table[1:]:
                rows.append([str(c or "").strip() for c in row])
            tables.append({"headers": headers, "rows": rows})
    except Exception:
        pass
    return tables


# ---------------------------------------------------------------------------
# DOCX extractor (section-aware, heading hierarchy)
# ---------------------------------------------------------------------------

def extract_docx_pages(file_bytes: bytes) -> list[RawPage]:
    """Split a Word document into logical sections at H1/H2 boundaries."""
    try:
        from docx import Document
    except ImportError:
        return []

    doc = Document(io.BytesIO(file_bytes))
    sections: list[RawPage] = []
    current_headings: list[str] = []  # breadcrumb stack
    current_texts: list[str] = []
    current_tables: list[dict] = []
    page_num = 1

    def _flush(headings: list[str], texts: list[str], tables: list[dict]) -> None:
        content = "\n\n".join(texts).strip()
        if content:
            sections.append(RawPage(
                page_number=len(sections) + 1,
                content=content,
                heading_breadcrumb=list(headings),
                section_path=".".join(str(i + 1) for i in range(len(headings))),
                tables=list(tables),
                is_toc=_looks_like_toc(content),
                is_bibliography=_looks_like_bibliography(content),
            ))

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if style.startswith("Heading"):
            # Extract heading level: "Heading 1" → 1, "Heading 2" → 2, etc.
            try:
                level = int(style.split()[-1])
            except (ValueError, IndexError):
                level = 1

            # Split at H1 or H2 boundaries
            if level <= 2 and current_texts:
                _flush(current_headings, current_texts, current_tables)
                current_texts = []
                current_tables = []

            # Update breadcrumb
            current_headings = current_headings[:level - 1] + [text]
        else:
            current_texts.append(text)

    # Extract tables
    for table in doc.tables:
        rows_data: list[list[str]] = []
        for row in table.rows:
            rows_data.append([cell.text.strip() for cell in row.cells])
        if rows_data:
            headers = rows_data[0]
            current_tables.append({"headers": headers, "rows": rows_data[1:]})

    # Flush last section
    if current_texts or current_tables:
        _flush(current_headings, current_texts, current_tables)

    return sections


# ---------------------------------------------------------------------------
# Markdown extractor
# ---------------------------------------------------------------------------

def extract_markdown_pages(file_bytes: bytes) -> list[RawPage]:
    """Split Markdown into sections at H2/H3 heading boundaries."""
    text = file_bytes.decode("utf-8", errors="replace")
    lines = text.splitlines()

    sections: list[RawPage] = []
    current_breadcrumb: list[str] = []
    current_lines: list[str] = []

    def _flush(breadcrumb: list[str], lines: list[str]) -> None:
        content = "\n".join(lines).strip()
        if content:
            sections.append(RawPage(
                page_number=len(sections) + 1,
                content=content,
                heading_breadcrumb=list(breadcrumb),
                section_path=".".join(str(i + 1) for i in range(len(breadcrumb))),
                tables=_extract_markdown_tables(content),
                is_toc=_looks_like_toc(content),
                is_bibliography=_looks_like_bibliography(content),
            ))

    for line in lines:
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            heading = m.group(2).strip()
            if level <= 3 and current_lines:
                _flush(current_breadcrumb, current_lines)
                current_lines = []
            current_breadcrumb = current_breadcrumb[:level - 1] + [heading]
        current_lines.append(line)

    if current_lines:
        _flush(current_breadcrumb, current_lines)

    return sections or [RawPage(page_number=1, content=text.strip())]


# ---------------------------------------------------------------------------
# HTML extractor
# ---------------------------------------------------------------------------

def extract_html_pages(file_bytes: bytes) -> list[RawPage]:
    """Split HTML into sections by heading or <section> / <article> tags."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        # Fallback: plain text extraction
        text = file_bytes.decode("utf-8", errors="replace")
        clean = re.sub(r"<[^>]+>", " ", text)
        return [RawPage(page_number=1, content=clean.strip())]

    soup = BeautifulSoup(file_bytes.decode("utf-8", errors="replace"), "html.parser")
    sections: list[RawPage] = []
    breadcrumb: list[str] = []

    def _add_section(element: Any, bcrumb: list[str]) -> None:
        tables = []
        for tbl in element.find_all("table"):
            headers = [th.get_text(strip=True) for th in tbl.find_all("th")]
            rows = []
            for tr in tbl.find_all("tr"):
                row = [td.get_text(strip=True) for td in tr.find_all("td")]
                if row:
                    rows.append(row)
            if rows:
                tables.append({"headers": headers, "rows": rows})

        text = element.get_text(separator="\n", strip=True)
        if text:
            sections.append(RawPage(
                page_number=len(sections) + 1,
                content=text,
                heading_breadcrumb=list(bcrumb),
                section_path=".".join(str(i + 1) for i in range(len(bcrumb))),
                tables=tables,
            ))

    # Try semantic sectioning first
    for section_tag in soup.find_all(["article", "section"]):
        heading = section_tag.find(re.compile(r"^h[1-6]$"))
        h_text = heading.get_text(strip=True) if heading else ""
        _add_section(section_tag, [h_text] if h_text else [])

    # Fall back: split at H2 headings
    if not sections:
        body = soup.find("body") or soup
        current: list[str] = []
        current_bcrumb: list[str] = []
        for elem in body.children:
            tag = getattr(elem, "name", None)
            if tag and re.match(r"^h[1-3]$", tag):
                if current:
                    _add_section(type("E", (), {"find_all": lambda *a, **k: [],
                                                "get_text": lambda **kw: "\n".join(current)})(),
                                 current_bcrumb)
                    current = []
                current_bcrumb = [elem.get_text(strip=True)]
            else:
                t = getattr(elem, "get_text", lambda **k: str(elem))(strip=True)
                if t:
                    current.append(t)
        if current:
            _add_section(type("E", (), {"find_all": lambda *a, **k: [],
                                        "get_text": lambda **kw: "\n".join(current)})(),
                         current_bcrumb)

    return sections or [RawPage(page_number=1,
                                content=soup.get_text(separator="\n", strip=True))]


# ---------------------------------------------------------------------------
# Plain text extractor
# ---------------------------------------------------------------------------

def extract_text_pages(file_bytes: bytes, chunk_chars: int = 3000) -> list[RawPage]:
    """Split plain text into fixed-size chunks with overlap."""
    text = file_bytes.decode("utf-8", errors="replace")
    overlap = 200
    pages: list[RawPage] = []
    start = 0
    page_num = 1
    while start < len(text):
        end = start + chunk_chars
        chunk = text[start:end].strip()
        if chunk:
            heading_breadcrumb = _detect_headings_from_text(chunk)
            pages.append(RawPage(
                page_number=page_num,
                content=chunk,
                heading_breadcrumb=heading_breadcrumb,
                section_path=str(page_num),
            ))
            page_num += 1
        start += chunk_chars - overlap
    return pages or [RawPage(page_number=1, content=text.strip())]


# ---------------------------------------------------------------------------
# Main dispatcher
# ---------------------------------------------------------------------------

def extract_pages(file_bytes: bytes, file_type: str) -> list[RawPage]:
    """Route to the appropriate extractor and return a list of RawPage objects."""
    ft = file_type.lower().lstrip(".")
    if ft == "pdf":
        return extract_pdf_pages(file_bytes)
    if ft in ("docx", "doc"):
        return extract_docx_pages(file_bytes)
    if ft in ("txt", "md", "markdown"):
        if ft == "md" or ft == "markdown":
            return extract_markdown_pages(file_bytes)
        return extract_text_pages(file_bytes)
    if ft in ("html", "htm"):
        return extract_html_pages(file_bytes)
    # CSV, JSON, XLSX — return a single page with a compact text representation
    try:
        text = file_bytes.decode("utf-8", errors="replace")[:10_000]
    except Exception:
        text = ""
    return [RawPage(page_number=1, content=text)] if text else []


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_TOC_PATTERNS   = re.compile(r"table of contents|contents\.{3,}|\.{5,}\d+", re.I)
_BIB_PATTERNS   = re.compile(r"bibliography|references\s*$|works cited", re.I)
_FIGURE_MARKERS = re.compile(r"figure\s+\d+|fig\.\s*\d+|exhibit\s+\d+", re.I)
_HEADING_RE     = re.compile(r"^(#{1,6})\s+(.+)$|^([A-Z][A-Z\s]{4,})$", re.M)


def _looks_like_toc(text: str) -> bool:
    return bool(_TOC_PATTERNS.search(text[:500]))


def _looks_like_bibliography(text: str) -> bool:
    return bool(_BIB_PATTERNS.search(text[:200]))


def _has_figure_markers(text: str) -> bool:
    return bool(_FIGURE_MARKERS.search(text))


def _detect_headings_from_text(text: str) -> list[str]:
    """Best-effort heading extraction from plain text (Markdown-style or ALLCAPS)."""
    headings = []
    for m in _HEADING_RE.finditer(text[:1000]):
        h = (m.group(2) or m.group(3) or "").strip()
        if h and len(h) < 120:
            headings.append(h)
        if len(headings) >= 3:
            break
    return headings
